import assemblyai as aai
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import time
import sys

aai.settings.api_key = ""

def transcribir_con_reintentos(audio_path, config, max_reintentos=3):
    transcriber = aai.Transcriber()
    for intento in range(max_reintentos):
        try:
            print(f"Intento {intento + 1} de transcripción...")
            transcript = transcriber.transcribe(audio_path, config)
            
            if transcript.status == aai.TranscriptStatus.error:
                raise Exception(f"Error en la transcripción: {transcript.error}")
                
            return transcript
            
        except Exception as e:
            print(f"Intento {intento + 1} fallido: {str(e)}")
            if intento < max_reintentos - 1:
                tiempo_espera = 5 * (intento + 1)
                print(f"Reintentando en {tiempo_espera} segundos...")
                time.sleep(tiempo_espera)
                continue
            raise

def ms_to_time(ms):
    seconds = int((ms / 1000) % 60)
    minutes = int((ms / (1000 * 60)) % 60)
    hours = int((ms / (1000 * 60 * 60)) % 24)
    return f"{hours:02}:{minutes:02}:{seconds:02}"

def crear_documento_word(transcript, output_file):
    try:
        doc = Document()
        doc.styles['Normal'].font.name = 'Arial'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

        doc.add_heading("Transcripción completa", level=1)
        doc.add_paragraph(transcript.text)

        doc.add_heading("Transcripción por bloques (Cada 2:45)", level=1)

        def add_styled_paragraph(doc, text, bold=False, underline=False, font_size=12):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(text)
            run.bold = bold
            run.underline = underline
            run.font.size = Pt(font_size)
            run.font.name = 'Arial'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

        # 2 minutos y 45 segundos = 165,000 milisegundos
        INTERVALO_MS = 165000 
        siguiente_marca = INTERVALO_MS
        
        texto_acumulado = []
        es_primera_marca = True

        # Primero filtramos todas las palabras que pertenecen SOLAMENTE al hablante A (Juan Gabriel Gomila)
        palabras_juan_gabriel = []
        for utterance in transcript.utterances:
            if utterance.speaker == "A":
                palabras_juan_gabriel.extend(utterance.words)

        if not palabras_juan_gabriel:
            print("No se encontraron intervenciones del hablante A.")
            return False

        # El primer bloque inicia en el tiempo de su primera palabra
        tiempo_inicio_bloque = palabras_juan_gabriel[0].start

        for word in palabras_juan_gabriel:
            # Si la palabra actual ya cruzó el límite de los 2:45, 5:30, etc.
            if word.start >= siguiente_marca:
                # 1. Soltamos el bloque que llevamos acumulado hasta ahora
                if texto_acumulado:
                    # Si es la primera marca de todo el documento, incluye el nombre del hablante
                    if es_primera_marca:
                        add_styled_paragraph(doc, "Juan Gabriel Gomila", bold=True, font_size=12)
                        es_primera_marca = False
                    
                    # Imprimimos la marca de tiempo de inicio de este bloque
                    marca_formateada = ms_to_time(tiempo_inicio_bloque)
                    add_styled_paragraph(doc, marca_formateada, bold=True, underline=True, font_size=12)
                    
                    # Imprimimos el texto acumulado
                    add_styled_paragraph(doc, " ".join(texto_acumulado), font_size=11)
                    doc.add_paragraph() # Espacio de separación

                # 2. Preparamos los datos para el NUEVO bloque
                texto_acumulado = [word.text]
                
                # Forzamos a que la línea de tiempo del nuevo bloque sea EXACTAMENTE el múltiplo de 2:45 correspondiente
                tiempo_inicio_bloque = siguiente_marca 
                
                # Avanzamos la siguiente marca en el tiempo (pueden ser varios intervalos si hubo un silencio largo)
                while word.start >= siguiente_marca:
                    siguiente_marca += INTERVALO_MS
            else:
                # Si no ha pasado el tiempo, seguimos acumulando palabras en el bloque actual
                texto_acumulado.append(word.text)

        # Imprimir el último bloque rezagado con el texto restante
        if texto_acumulado:
            if es_primera_marca:
                add_styled_paragraph(doc, "Juan Gabriel Gomila", bold=True, font_size=12)
            marca_formateada = ms_to_time(tiempo_inicio_bloque)
            add_styled_paragraph(doc, marca_formateada, bold=True, underline=True, font_size=12)
            add_styled_paragraph(doc, " ".join(texto_acumulado), font_size=11)

        doc.save(output_file)
        print(f"Documento guardado exitosamente en: {output_file}")
        return True
        
    except Exception as e:
        print(f"Error al crear el documento Word: {str(e)}")
        return False

def main():
    audio_file = "202.mp3"
    output_file = "Frogmación 200: “Me fui a la India… y lo que descubrí me cambió para siempre”.docx"

    config = aai.TranscriptionConfig(
        speaker_labels=True,
        language_code="es",
    )

    try:
        # Paso 1: Transcribir el audio con reintentos
        transcript = transcribir_con_reintentos(audio_file, config)
        
        # Paso 2: Crear el documento Word
        if not crear_documento_word(transcript, output_file):
            sys.exit(1)
            
    except Exception as e:
        print(f"Error crítico: {str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    main()
