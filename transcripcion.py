import assemblyai as aai
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import time
import sys

aai.settings.api_key = "46f568c434a349df894e5f74829619d2"

def transcribir_con_reintentos(audio_path, config, max_reintentos=3):
    transcriber = aai.Transcriber()
    for intento in range(max_reintentos):
        try:
            print(f"Intento {intento + 1} de transcripción...")
            transcript = transcriber.transcribe(audio_path, config)
            
            if transcript.status == aai.TranscriptStatus.error:
                raise Exception(f"Error en la transcripción: {transcript.error}")
                
            return transcript
            
        except Exception as e:
            print(f"Intento {intento + 1} fallido: {str(e)}")
            if intento < max_reintentos - 1:
                tiempo_espera = 5 * (intento + 1)  # Espera progresiva: 5, 10, 15 segundos
                print(f"Reintentando en {tiempo_espera} segundos...")
                time.sleep(tiempo_espera)
                continue
            raise  # Si todos los reintentos fallan

def ms_to_time(ms):
    seconds = int((ms / 1000) % 60)
    minutes = int((ms / (1000 * 60)) % 60)
    hours = int((ms / (1000 * 60 * 60)) % 24)
    return f"{hours:02}:{minutes:02}:{seconds:02}"

def crear_documento_word(transcript, output_file):
    try:
        doc = Document()
        doc.styles['Normal'].font.name = 'Arial'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

        doc.add_heading("Transcripción completa", level=1)
        doc.add_paragraph(transcript.text)

        doc.add_heading("Transcripción por hablantes", level=1)

        def add_styled_paragraph(doc, text, bold=False, underline=False, font_size=12):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(text)
            run.bold = bold
            run.underline = underline
            run.font.size = Pt(font_size)
            run.font.name = 'Arial'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

        for utterance in transcript.utterances:
            if utterance.speaker == "A":
                speaker_name = "Juan Gabriel Gomila"
            elif utterance.speaker == "B":
                speaker_name = "Sebastián Barajas Caseny"
            else:
                speaker_name = f"Speaker {utterance.speaker}"

            start_time = ms_to_time(utterance.start)

            add_styled_paragraph(doc, speaker_name, bold=True, font_size=12)
            add_styled_paragraph(doc, start_time, bold=True, underline=True, font_size=12)
            add_styled_paragraph(doc, utterance.text, font_size=11)
            doc.add_paragraph()

        doc.save(output_file)
        print(f"Documento guardado exitosamente en: {output_file}")
        return True
        
    except Exception as e:
        print(f"Error al crear el documento Word: {str(e)}")
        return False

def main():
    audio_file = "C:/Users/Skywalker/Documents/Programacion/Transcriptor/Transcriptor/Frogmacion158.mp3"
    output_file = "transcripcion.docx"

    config = aai.TranscriptionConfig(
        speaker_labels=True,
        language_code="es",
    )

    try:
        # Paso 1: Transcribir el audio con reintentos
        transcript = transcribir_con_reintentos(audio_file, config)
        
        # Paso 2: Crear el documento Word
        if not crear_documento_word(transcript, output_file):
            sys.exit(1)
            
    except Exception as e:
        print(f"Error crítico: {str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    main()